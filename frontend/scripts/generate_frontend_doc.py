"""Generate OpenAssess Frontend Process documentation as Word (.docx)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

OUTPUT = Path(__file__).resolve().parents[1] / "docs" / "OpenAssess_Frontend_Process.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def build_document():
    doc = Document()

    title = doc.add_heading("OpenAssess Frontend Process", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Technical Documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x0D, 0x94, 0x88)

    doc.add_paragraph("Version: MVP (Demo Mode) | Stack: Next.js 16, React 19, Tailwind CSS 4")
    doc.add_paragraph()

    # 1. Overview
    add_heading(doc, "1. Project Overview", 1)
    add_para(
        doc,
        "OpenAssess is an AI-powered continuous assessment platform. The frontend allows students "
        "to sign in, view progress on a dashboard, select assessment topics, take timed quizzes, "
        "review AI-powered results, and browse a verified knowledge portfolio. The current build "
        "is an MVP that uses mock data on the client side; it is designed so real FastAPI backend "
        "APIs can be connected later without restructuring the UI.",
    )

    add_heading(doc, "1.1 Core Philosophy", 2)
    add_para(doc, "Learn → Assess → Improve → Retry → Mastery", bold=True)

    add_heading(doc, "1.2 Technology Stack", 2)
    for item in [
        "Next.js 16 (App Router) — routing, layouts, server and client components",
        "React 19 — UI components and interactive quiz logic",
        "TypeScript — type-safe code across pages and shared libraries",
        "Tailwind CSS 4 — utility-first styling and responsive design",
        "CSS Modules — isolated animations for AnimatedBackground",
        "Geist font — typography via next/font",
    ]:
        add_bullet(doc, item)

    # 2. Folder structure
    add_heading(doc, "2. Project Structure", 1)
    add_para(doc, "The frontend lives in the frontend/ directory with this layout:")

    structure = """frontend/
├── app/                          # Next.js App Router pages
│   ├── layout.tsx                # Root layout (fonts, metadata)
│   ├── page.tsx                  # Login page (home route /)
│   ├── globals.css               # Global styles and animations
│   └── dashboard/                # Protected student area
│       ├── layout.tsx            # Sidebar + main content shell
│       ├── page.tsx              # Dashboard home
│       ├── assessment/
│       │   ├── page.tsx          # Topic selection
│       │   ├── take/page.tsx     # Live quiz (client component)
│       │   └── results/page.tsx  # Score and AI insights
│       └── portfolio/page.tsx    # Certifications and heatmap
├── components/                   # Reusable UI
│   ├── AnimatedBackground.tsx    # Full-screen animated backdrop
│   ├── animated-background.module.css
│   ├── ui.tsx                    # Card, buttons, badges, progress bars
│   ├── icons.tsx                 # SVG icon components
│   ├── brand-logo.tsx            # OpenAssess branding
│   └── heatmap.tsx               # Study activity heatmap grid
├── lib/data.ts                   # Mock data (topics, questions, stats)
└── package.json                  # Dependencies and npm scripts"""

    p = doc.add_paragraph()
    run = p.add_run(structure)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # 3. User flow
    add_heading(doc, "3. End-to-End User Flow", 1)
    add_para(
        doc,
        "The MVP supports a complete student journey from login to portfolio review. "
        "All navigation uses Next.js Link and the App Router; there is no real authentication yet—"
        "the Sign In button routes directly to the dashboard.",
    )

    flows = [
        ("Step 1 — Login (/)", "User lands on the login page with AnimatedBackground and a frosted-glass sign-in card. Email and password fields are displayed; submitting navigates to /dashboard (demo mode)."),
        ("Step 2 — Dashboard (/dashboard)", "Shows welcome message, stat cards (quizzes, mastery, hours), topic mastery progress bars, recent activity feed, and a study consistency heatmap preview."),
        ("Step 3 — Topic Selection (/dashboard/assessment)", "User picks a subject (Mathematics, Computer Science, History, Physics). Each card shows difficulty, question count, and mastery percentage. Start Assessment links to the quiz."),
        ("Step 4 — Take Quiz (/dashboard/assessment/take)", "Client-side quiz with 3 mock Computer Science questions. Features: 15-minute countdown timer, question navigator, progress bar, MCQ selection, Previous/Next/Submit. Answers stored in React state."),
        ("Step 5 — Results (/dashboard/assessment/results?score=N)", "After submit or timeout, user is redirected with score in URL query. Page shows percentage, strengths, areas for review, AI insight summary, and per-question explanations."),
        ("Step 6 — Portfolio (/dashboard/portfolio)", "Displays learner profile, micro-certifications, topic mastery chart, and full 12-week activity heatmap."),
    ]
    for title, desc in flows:
        add_heading(doc, title, 2)
        add_para(doc, desc)

    # 4. Routing
    add_heading(doc, "4. Routing and Layouts", 1)
    add_para(
        doc,
        "Next.js App Router maps folders to URLs. layout.tsx files wrap child pages and persist "
        "across navigation within the same segment.",
    )
    add_bullet(doc, "app/layout.tsx — Applies to every page; loads Geist fonts and global CSS.")
    add_bullet(doc, "app/dashboard/layout.tsx — Adds the sidebar and main content area for all dashboard routes.")
    add_bullet(doc, "Client components use \"use client\" when they need hooks (useState, useRouter, useSearchParams).")

    # 5. Components
    add_heading(doc, "5. Key Components Explained", 1)

    components = [
        ("AnimatedBackground", "Fixed full-viewport layer (z-index 0) with navy base (#1A2642), four drifting teal/sky blurred orbs, dot-grid overlay, and CSS keyframe animations. Does not block clicks or affect scroll."),
        ("Dashboard Sidebar", "Responsive navigation with Dashboard, Assessments, and Portfolio links. Highlights active route. Mobile hamburger menu with overlay."),
        ("UI Kit (components/ui.tsx)", "Shared Card, PageHeader, Badge, ProgressBar, StatCard, and ButtonLink for consistent dashboard styling."),
        ("ProgressHeatmap", "Renders a GitHub-style grid from HEATMAP_DATA in lib/data.ts using intensity-colored cells."),
        ("lib/data.ts", "Central mock data store. When the backend is ready, replace or supplement this file with API fetch functions."),
    ]
    for name, desc in components:
        add_heading(doc, name, 2)
        add_para(doc, desc)

    # 6. Quiz logic
    add_heading(doc, "6. Quiz Engine (Client-Side)", 1)
    add_para(
        doc,
        "The quiz page (take/page.tsx) is a Client Component because it requires timers and user interaction.",
    )
    for item in [
        "MOCK_QUESTIONS array defines question text, four options, correct answer index, and AI explanation.",
        "selectedAnswers state tracks one answer per question (null if unanswered).",
        "Timer counts down from 900 seconds (15 minutes); at zero, quiz auto-submits.",
        "calculateScorePercent compares answers to correctAnswer indices and rounds to a percentage.",
        "router.push navigates to /dashboard/assessment/results?score={percent}.",
        "Results page reads score via useSearchParams (wrapped in Suspense for Next.js compatibility).",
    ]:
        add_bullet(doc, item)

    # 7. Styling
    add_heading(doc, "7. Styling Approach", 1)
    add_bullet(doc, "Tailwind utility classes for layout, spacing, colors, and responsive breakpoints (sm, lg).")
    add_bullet(doc, "globals.css — Tailwind import, custom keyframes (login-card-enter, fade-in-up).")
    add_bullet(doc, "Login page — Frosted glass (backdrop-blur, bg-white/10), white text, teal primary button (#0D9488) with hover glow.")
    add_bullet(doc, "Dashboard — Light slate background with indigo accent brand colors.")

    # 8. Running
    add_heading(doc, "8. How to Run the Frontend", 1)
    steps = [
        "Open a terminal in the frontend/ directory.",
        "Run: npm install (first time only).",
        "Run: npm run dev — starts development server at http://localhost:3000.",
        "Run: npm run build — production build.",
        "Run: npm start — serve production build.",
    ]
    for i, step in enumerate(steps, 1):
        add_para(doc, f"{i}. {step}")

    # 9. Backend integration
    add_heading(doc, "9. Future Backend Integration", 1)
    add_para(
        doc,
        "Developer A (Backend) will provide FastAPI endpoints. The frontend will connect using fetch or Axios:",
    )
    apis = [
        "POST /auth/login — replace demo sign-in",
        "GET /questions?topic=... — load quiz questions",
        "POST /quiz/submit — send answers, receive score and AI explanations",
        "GET /analytics — dashboard stats and heatmap data",
        "GET /certifications — portfolio credentials",
    ]
    for api in apis:
        add_bullet(doc, api)
    add_para(
        doc,
        "Recommended pattern: create frontend/services/api.ts with typed fetch wrappers; "
        "keep pages thin and move data fetching to Server Components or React hooks as appropriate.",
    )

    # 10. Deployment
    add_heading(doc, "10. Deployment", 1)
    add_para(
        doc,
        "Per project brief, the frontend targets Vercel. Connect the GitHub repository, "
        "set the root directory to frontend/, and deploy. Environment variables (e.g. NEXT_PUBLIC_API_URL) "
        "will point to the FastAPI backend on Render or Railway.",
    )

    # Diagram as text
    add_heading(doc, "11. Architecture Diagram (Text)", 1)
    diagram = (
        "Browser\n"
        "   |\n"
        "   v\n"
        "Next.js App Router (pages + layouts)\n"
        "   |\n"
        "   +-- Server Components (static pages, metadata)\n"
        "   |\n"
        "   +-- Client Components (quiz, sidebar, results)\n"
        "   |\n"
        "   v\n"
        "components/ + lib/data.ts (UI + mock data)\n"
        "   |\n"
        "   v  [Future]\n"
        "FastAPI Backend + PostgreSQL + AI APIs"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    doc.add_paragraph()
    add_para(
        doc,
        "Document generated for OpenAssess frontend (Developer B — UI). "
        "For full project roles and roadmap, see project_brief.md in the repository.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"Created: {path}")
