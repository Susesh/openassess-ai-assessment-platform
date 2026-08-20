"""Generate OpenAssess backend explanation as a Word document."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

OUTPUT = r"c:\Users\karth\OneDrive\Desktop\Vortex_iq\Openassess\OpenAssess_Backend_Explanation.docx"


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()

    title = doc.add_heading("OpenAssess Backend — Explanation Document", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "Version 1.0.0  |  FastAPI + PostgreSQL + SQLAlchemy + Google Gemini",
    )
    doc.add_paragraph()

    # --- What is it ---
    add_heading(doc, "1. What Is This Backend?", 1)
    add_para(
        doc,
        "OpenAssess is a REST API (web service) for an AI-powered learning and assessment "
        "platform. It does not include a frontend in this repository — it is the server that "
        "a web or mobile application talks to over HTTP.",
    )
    add_para(doc, "One-sentence summary:", bold=True)
    add_para(
        doc,
        '"OpenAssess is a FastAPI backend that lets students register, take topic-based quizzes, '
        "get AI explanations when they're wrong, track their progress, earn certificates, and "
        'run proctored exams — all backed by PostgreSQL."',
    )

    # --- Problem ---
    add_heading(doc, "2. What Problem Does It Solve?", 1)
    add_para(
        doc,
        "Traditional tests are often one-shot (single exam, little feedback). OpenAssess "
        "supports continuous assessment:",
    )
    for item in [
        "Many attempts over time",
        "Feedback after mistakes (including Google Gemini explanations)",
        "Analytics on strengths and weaknesses",
        "Certificates when performance is strong enough",
        "Proctoring logs for exam integrity",
    ]:
        add_bullet(doc, item)

    # --- Tech stack ---
    add_heading(doc, "3. Technology Stack", 1)
    add_table(
        doc,
        ["Technology", "Role"],
        [
            ["FastAPI", "HTTP API framework; auto documentation at /docs"],
            ["PostgreSQL", "Database (users, questions, scores, certificates, logs)"],
            ["SQLAlchemy", "Object-relational mapping (ORM) for database access"],
            ["Pydantic", "Validates request/response JSON (schemas/ folder)"],
            ["JWT + bcrypt", "Secure login and password hashing"],
            ["Google Gemini", "AI explanations for incorrect quiz answers"],
        ],
    )

    # --- Architecture ---
    add_heading(doc, "4. Code Organization (Layers)", 1)
    add_para(doc, "The backend follows a layered architecture:")
    layers = [
        "Client (future frontend) → HTTP requests",
        "main.py → Application entry, CORS, /health, router registration",
        "routes/ → API endpoints (URL handlers)",
        "schemas/ → JSON request/response shapes (Swagger documentation)",
        "services/ → Business logic helpers (AI, certifications)",
        "utils/ → Authentication utilities (JWT, passwords)",
        "models/ → Database table definitions",
        "database.py → PostgreSQL connection and sessions",
    ]
    for layer in layers:
        add_bullet(doc, layer)

    # --- API groups ---
    add_heading(doc, "5. Main API Features", 1)

    add_heading(doc, "5.1 Authentication (Auth)", 2)
    add_para(doc, "Register, login, and view profile (/auth/register, /auth/login, /auth/me).")
    add_para(doc, "Protected routes require: Authorization: Bearer <token>")

    add_heading(doc, "5.2 Questions", 2)
    add_bullet(doc, "List topics and subtopics (curriculum structure)")
    add_bullet(doc, "Get random questions for practice (filtered by topic/subtopic)")
    add_bullet(doc, "Admin users can add new questions (ADMIN_EMAILS in .env)")

    add_heading(doc, "5.3 Quiz (Core Flow)", 2)
    add_bullet(doc, "POST /quiz/start — Select random questions, create an Attempt, return questions without correct answers")
    add_bullet(doc, "POST /quiz/submit — Grade answers, save Results, update score; Gemini provides ai_explanation for wrong answers")

    add_heading(doc, "5.4 Analytics", 2)
    add_bullet(doc, "GET /analytics/me — Total attempts, average score, pass rate, strongest/weakest topic")
    add_bullet(doc, "GET /analytics/heatmap — Per-topic performance data for charts")

    add_heading(doc, "5.5 Certifications", 2)
    add_bullet(doc, "POST /certifications/generate — Issue certificate when topic average score ≥ 80%")
    add_bullet(doc, "GET /certifications/me — List all certificates for the logged-in user")

    add_heading(doc, "5.6 Proctoring", 2)
    add_bullet(doc, "POST /proctoring/log — Record integrity events (tab switch, face not detected, multiple faces)")
    add_bullet(doc, "GET /proctoring/report/{attempt_id} — Review all events for a quiz attempt")

    add_heading(doc, "5.7 Health", 2)
    add_bullet(doc, "GET /health — API status, version, and database connectivity")

    # --- Data model ---
    add_heading(doc, "6. Database Entities", 1)
    add_table(
        doc,
        ["Entity", "Description"],
        [
            ["User", "Student account (name, email, hashed password)"],
            ["Topic / Subtopic", "Subject structure (e.g., Python → Functions)"],
            ["Question", "Multiple-choice question with 4 options and explanation"],
            ["Attempt", "One quiz session (started → submitted)"],
            ["Result", "Single answer within an attempt"],
            ["Certification", "Earned certificate for a topic (unique certificate code)"],
            ["ProctorLog", "Integrity event logged during an attempt"],
        ],
    )

    # --- Student journey ---
    add_heading(doc, "7. Student Journey (Demo Flow)", 1)
    steps = [
        "Register or log in",
        "Browse topics and subtopics",
        "Start a quiz (POST /quiz/start)",
        "Answer questions in the frontend",
        "Submit answers (POST /quiz/submit)",
        "Receive score, pass/fail, and AI explanations for wrong answers",
        "View analytics dashboard (GET /analytics/me, /analytics/heatmap)",
        "Generate certificate if average score ≥ 80% (POST /certifications/generate)",
    ]
    for i, step in enumerate(steps, 1):
        add_bullet(doc, f"Step {i}: {step}")

    # --- Presentation scripts ---
    add_heading(doc, "8. How to Explain (Presentation Scripts)", 1)

    add_heading(doc, "8.1 Non-Technical Audience (30 seconds)", 2)
    add_para(
        doc,
        '"It is the brain of our assessment app. Students log in, take quizzes by subject, '
        "get instant AI help when they miss a question, see how they're improving over time, "
        "and can earn certificates. Admins can add questions, and the system can log cheating "
        'signals like tab switches."',
    )

    add_heading(doc, "8.2 Technical Audience (1 minute)", 2)
    add_para(
        doc,
        '"Layered FastAPI service with SQLAlchemy on PostgreSQL. JWT authentication, Pydantic '
        "schemas for OpenAPI at /docs, modular routers per domain. Quiz flow separates start "
        "and submit with attempt-bound question IDs. Wrong-answer feedback uses async Gemini "
        "calls via services/ai_service.py. Analytics aggregate completed attempts; "
        'certifications require 80% topic average."',
    )

    add_heading(doc, "8.3 What is /docs?", 2)
    add_para(
        doc,
        '"Interactive API documentation — you can try every endpoint from the browser without '
        'building a frontend first."',
    )

    # --- Limitations ---
    add_heading(doc, "9. Current Scope and Notes", 1)
    add_bullet(doc, "This repository contains the API only — no React/Vue frontend included")
    add_bullet(doc, "Run seed.py once to populate sample Python and SQL questions for demos")
    add_bullet(doc, "Interactive documentation available at http://127.0.0.1:8000/docs when the server is running")
    add_bullet(doc, "Start server: uvicorn main:app --reload (from the backend folder)")

    doc.add_paragraph()
    footer = doc.add_paragraph("OpenAssess — AI-Powered Continuous Assessment Platform")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build()
